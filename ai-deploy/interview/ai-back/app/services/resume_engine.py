"""
简历分析核心引擎
负责：Word 解析、文本分块、LLM 逐块分析、整体评价、输出 Word 生成
"""
import json
import os
import re
from pathlib import Path
from typing import List, Dict

from docx import Document
from langchain_core.messages import HumanMessage, SystemMessage

UPLOAD_DIR = Path("resume_uploads")
OUTPUT_DIR = Path("resume_outputs")

# 启动时确保目录存在
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)


# ============================================================
# 1. 解析 DOCX → 有意义的"块"列表
# ============================================================
def parse_docx_to_blocks(file_path: str) -> List[Dict]:
    """
    解析 Word 文档，不仅遍历段落，还遍历表格内容。
    按章节标题或逻辑间隙分组。
    """
    doc = Document(file_path)
    blocks = []
    current_section = "简历内容"
    current_paragraphs = []

    def flush_block():
        """将当前积累的段落刷入块列表"""
        text = "\n".join(p.strip() for p in current_paragraphs if p.strip())
        if text:
            blocks.append({
                "block_index": len(blocks),
                "section_title": current_section,
                "original_text": text,
            })

    # 先处理所有普通段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 识别章节标题：Heading 样式，或者全加粗短行
        is_heading = (
            para.style.name.startswith("Heading")
            or (len(text) < 30 and para.runs and all(r.bold for r in para.runs if r.text.strip()))
        )

        if is_heading:
            flush_block()
            current_section = text
            current_paragraphs = []
        else:
            current_paragraphs.append(text)

    # 再处理所有表格中的内容（很多简历是表格布局）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        current_paragraphs.append(para.text.strip())
    
    flush_block()  # 清空最后的积累

    # 如果完全没读到内容，报个错，防止后续 LLM 报错
    if not blocks:
        blocks.append({
            "block_index": 0,
            "section_title": "基础信息",
            "original_text": "未在此 Word 文档中识别到明显文字，请检查文件内容或尝试非表格排版的简历。",
        })
        
    return blocks


# ============================================================
# 2. 单块 LLM 分析
# ============================================================
def analyze_single_block(block_text: str, section_title: str, job_target: str, llm) -> Dict:
    """
    调用 LLM 对单个文本块进行分析，返回修改建议和原因。
    要求 LLM 以 JSON 格式回答，保证结构稳定性。
    """
    system_prompt = (
        "你是一位拥有 10 年以上经验的顶级职业规划导师和简历专家。"
        "你的任务是对候选人简历中的某个段落提出针对性的专业修改建议，使其更能打动 HR。\n"
        "【特别要求】：如果原始文字非常简短单薄（少于20字），请大胆进行引导式扩充，补充符合该岗位常见的高级技术细节、关键动作与成果描述；如果原文字数已经详细，则专注提炼核心价值并修饰专业素养。不用考虑字数增多破坏排版的问题。\n"
        "你必须以严格的 JSON 格式回答，不要输出 JSON 以外的任何内容。"
        '格式为: {"suggested": "修改后的完整文本", "reason": "修改原因（1-2句话，简洁专业）"}'
        "如果该段落已经足够优秀无需修改，suggested 字段返回原文，reason 字段说明'此段表述清晰，无需大幅修改'。"
    )

    user_prompt = (
        f"候选人的目标岗位：{job_target}\n"
        f"当前章节：{section_title}\n\n"
        f"待分析的简历内容：\n{block_text}\n\n"
        "请给出修改建议（严格 JSON 格式）："
    )

    try:
        messages = [SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)]
        response = llm.invoke(messages)
        raw = response.content.strip()

        # 强化抽取：只保留第一个 { 到最后一个 } 之间的内容
        start_idx = raw.find('{')
        end_idx = raw.rfind('}')
        if start_idx != -1 and end_idx != -1:
            raw = raw[start_idx:end_idx+1]

        # strict=False 允许字符串内存在真实的回车\换行符
        result = json.loads(raw, strict=False)
        return {
            "suggested_text": result.get("suggested", block_text),
            "suggestion_reason": result.get("reason", ""),
        }
    except Exception as e:
        # 分析本块失败时，优雅降级：保留原文，标注失败
        print(f"[ResumeEngine] 块分析失败 (原文截取:{raw[:100]}...): {e}")
        return {
            "suggested_text": block_text,
            "suggestion_reason": f"💡 AI 分析此块时遇到异常，建议人工复查。",
        }


# ============================================================
# 3. 整体评价
# ============================================================
def generate_overall_evaluation(all_blocks: List[Dict], job_target: str, llm) -> Dict:
    """
    基于所有块的原始内容，生成整体优点、缺点、岗位匹配度评分。
    """
    full_resume = "\n\n".join(
        f"【{b['section_title']}】\n{b['original_text']}" for b in all_blocks
    )

    system_prompt = (
        "你是一位顶级简历评估专家。请对以下完整简历进行综合诊断，"
        "以严格 JSON 格式返回分析结果，不要输出 JSON 以外的任何内容。\n"
        '格式为: {"strengths": ["优点1", "优点2", ...], "weaknesses": ["缺点1", "缺点2", ...], "job_match_score": 85.0}'
        "\njob_match_score 为 0-100 的浮点数，代表该简历与目标岗位的契合度。"
    )

    user_prompt = f"目标岗位：{job_target}\n\n完整简历内容：\n{full_resume}"

    try:
        messages = [SystemMessage(content=system_prompt), HumanMessage(content=user_prompt)]
        response = llm.invoke(messages)
        raw = response.content.strip()
        
        # 强化抽取
        start_idx = raw.find('{')
        end_idx = raw.rfind('}')
        if start_idx != -1 and end_idx != -1:
            raw = raw[start_idx:end_idx+1]

        result = json.loads(raw, strict=False)
        return {
            "overall_strengths": json.dumps(result.get("strengths", []), ensure_ascii=False),
            "overall_weaknesses": json.dumps(result.get("weaknesses", []), ensure_ascii=False),
            "job_match_score": float(result.get("job_match_score", 0.0)),
        }
    except Exception as e:
        print(f"[ResumeEngine] 整体评价失败 (原文截取:{raw[:100]}...): {e}")
        return {
            "overall_strengths": "[]",
            "overall_weaknesses": "[]",
            "job_match_score": 0.0,
        }


# ============================================================
# 4. 构建最终输出 Word (双模式)
# ============================================================
def build_output_docx(analysis_id: int, original_file_path: str, blocks_data: List[Dict]) -> Dict[str, str]:
    """
    根据用户审阅结果，在硬盘上同时生成原模板修改版和纯净内容版。
    返回 { "full": 完整版路径, "content": 纯净版路径 }
    """
    from docx.oxml.ns import qn
    from docx.shared import Pt

    # ---------------- 模式 1：生成完整版（保留原排版） ----------------
    doc = Document(original_file_path)
    
    block_index = 0
    current_paragraphs = []
    block_map = {}

    def flush_block():
        nonlocal block_index, current_paragraphs
        valid_paras = [p for p in current_paragraphs if p.text.strip()]
        if valid_paras:
            block_map[block_index] = valid_paras
            block_index += 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        is_heading = (para.style.name.startswith("Heading") or (len(text) < 30 and para.runs and all(r.bold for r in para.runs if r.text.strip())))
        if is_heading:
            flush_block()
            current_paragraphs = []
        else:
            current_paragraphs.append(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        current_paragraphs.append(para)
    flush_block()

    for block in blocks_data:
        idx = block["block_index"]
        if block["is_accepted"] is True and idx in block_map:
            suggested = block["suggested_text"] or block["original_text"]
            paras = block_map[idx]
            if len(paras) > 0:
                paras[0].text = suggested
                for other_para in paras[1:]:
                    other_para.text = ""

    output_full_path = OUTPUT_DIR / f"resume_output_{analysis_id}_full.docx"
    doc.save(str(output_full_path))

    # ---------------- 模式 2：生成纯净内容版 ----------------
    new_doc = Document()
    style = new_doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for i in range(1, 10):
        h_style_name = f'Heading {i}'
        if h_style_name in new_doc.styles:
            h_style = new_doc.styles[h_style_name]
            h_style.font.name = '宋体'
            h_style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            h_style.font.color.rgb = None

    new_doc.core_properties.author = "AI 简历助手"
    current_section = None
    for block in sorted(blocks_data, key=lambda x: x["block_index"]):
        if block["section_title"] != current_section:
            current_section = block["section_title"]
            new_doc.add_heading(current_section, level=1)
            
        final_text = block["suggested_text"] if block["is_accepted"] else block["original_text"]
        for line in final_text.split("\n"):
            if line.strip():
                new_doc.add_paragraph(line.strip())

    output_content_path = OUTPUT_DIR / f"resume_output_{analysis_id}_content.docx"
    new_doc.save(str(output_content_path))

    return {
        "full": str(output_full_path),
        "content": str(output_content_path)
    }
